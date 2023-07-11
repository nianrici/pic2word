import os
import math
from docx import Document
from docx.shared import Cm

# Creamos una instancia de Document() para crear un documento de Word
document = Document()

# Saltamos una página y creamos una tabla con dos columnas y una fila para el título
document.add_page_break()
table = document.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False

# Establecemos el ancho de las columnas en centímetros
table.columns[0].width = Cm(10)
table.columns[1].width = Cm(10)

# Especificamos la carpeta de la que queremos extraer las imágenes
my_dir = "."

# Contamos la cantidad de imágenes que se insertarán
piccount = 0
for root, dirs, files in os.walk(my_dir):
    for filename in files:
        if filename.upper().endswith(('.PNG', '.JPG')):
            piccount += 1
print(piccount, " imágenes serán insertadas")

# Especificamos el número de columnas que deseamos en la tabla
total_column = 2

# Calculamos el número total de filas necesarias para acomodar todas las imágenes
total_row = math.ceil(piccount / total_column) + 2

# Añadimos las filas necesarias a la tabla
for row in range(total_row):
    table.add_row()

# Reiniciamos el contador de imágenes
piccount = 0

# Iteramos a través de los archivos en la carpeta y sus subcarpetas
for root, dirs, files in os.walk(my_dir):
    for filename in files:
        if filename.upper().endswith(('.PNG', '.JPG')):
            print(filename)

            # Calculamos el número de columna y fila para la celda actual
            cell_column = (piccount % total_column)
            cell_row = (piccount // total_column) + 2
            print('cell_column={}, cell_row={}'.format(cell_column, cell_row))

            # Seleccionamos la celda actual
            cell = table.cell(cell_row, cell_column)

            # Añadimos la imagen a la celda
            current_pic = cell.add_paragraph().add_run()
            current_pic.add_picture(os.path.join(root, filename), width=Cm(10))

            # Añadimos el nombre del archivo sin la extensión como pie de foto
            cell.add_paragraph(os.path.splitext(filename)[0], style='Caption')
            piccount += 1

# Guardamos el documento
document.save('images.docx')
